from docx import Document
from funcoes_requisicao import processar_texto
import os
import pytesseract
from PIL import Image

def transcrever_imagem(caminho_imagem):
    """
    Transcreve o texto de uma imagem usando OCR.
    
    :param caminho_imagem: Caminho da imagem a ser processada.
    :return: Texto extraído da imagem.
    """
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    imagem = Image.open(caminho_imagem)
    texto_extraido = pytesseract.image_to_string(imagem)  # 'por' para Português
    return texto_extraido

def adicionar_run(paragraph, texto, negrito=False):
    """
    Adiciona um run ao parágrafo com a formatação especificada.
    
    :param paragraph: Objeto Paragraph do python-docx.
    :param texto: Texto a ser adicionado ao run.
    :param negrito: Indica se o texto deve ser negrito.
    """
    run = paragraph.add_run(texto)
    if negrito:
        run.bold = True  # Aplica negrito se necessário

def substituir_placeholders_no_paragrafo(paragraph, informacoes):
    """
    Substitui os placeholders no parágrafo com os valores fornecidos.
    
    :param paragraph: Objeto Paragraph do python-docx.
    :param informacoes: Dicionário com chaves correspondentes aos placeholders e seus valores.
    """
    # Limpa o parágrafo existente
    paragraph.clear()  # Limpa todos os runs existentes

    # Adiciona o texto formatado com negrito onde necessário
    adicionar_run(paragraph, "Em virtude ao atendimento a", negrito=False)
    adicionar_run(paragraph, " REQUISIÇÃO ONLINE DE PERICIA Nº ", negrito=True)

    adicionar_run(paragraph, informacoes['requisicao'], negrito=True)  # Negrito para requisição
    adicionar_run(paragraph, ", datada de ", negrito=False)
    adicionar_run(paragraph, informacoes['data'], negrito=True)  # Negrito para data
    adicionar_run(paragraph, ", referente ao, ", negrito=False)
    adicionar_run(paragraph, "INQUÉRITO POR PORTARIA Nº ", negrito=True)

    adicionar_run(paragraph, informacoes['inquerito'], negrito=True)  # Negrito para inquérito
    adicionar_run(paragraph, " – DIVISÃO DE COMBATE A CRIMES CONTRA DIREITOS INDIVIDUAIS POR MEIOS CIBERNÉTICOS, ", negrito=False)
    adicionar_run(paragraph, "e assinado pela autoridade acima mencionada, solicitando", negrito=True)
    adicionar_run(paragraph, " perícia em aparelho celular", negrito=False)
    adicionar_run(paragraph, " a fim de", negrito=False)
    adicionar_run(paragraph, " extração de dados (registro de chamadas, contatos, fotos, imagens, áudios, vídeos, conversas de aplicativos e de mensagens de texto) e análise de conteúdo", negrito=True)
    adicionar_run(paragraph, ", a fim de colaborar com as investigações. O aparelho de telefonia celular foi recebido pelo perito signatário para exame pericial onde se constatou que o aparelho encontrava-se", negrito=False)
    adicionar_run(paragraph, " lacrado (Lacre nº ",negrito=True)
    adicionar_run(paragraph, informacoes['lacre'], negrito=True)  # Negrito para lacre
    adicionar_run(paragraph, ") no saco de evidências ", negrito=False)
    adicionar_run(paragraph, "(Saco nº A231650563)", negrito=False)
    adicionar_run(paragraph, "(ver  ", negrito=False)
    adicionar_run(paragraph, " 'Ilustração 01 e 02') ", negrito=False)
    adicionar_run(paragraph, "em seguida foi deslacrado pelo Perito ", negrito=False)
    adicionar_run(paragraph, "(ver Ilustração 03) ", negrito=False)
    adicionar_run(paragraph, ". Após ligar o aparelho, o Perito observou que o aparelho de telefonia celular estava em , conforme   ", negrito=False)
    adicionar_run(paragraph, "modo avião ", negrito=True)
    adicionar_run(paragraph, "“Ilustração 07” ", negrito=True)

def salvar_documento_com_tratamento(doc, caminho_inicial):
    """
    Salva o documento, tratando erros de permissão adicionando um número ao nome do arquivo se necessário.
    
    :param doc: Objeto Document do python-docx.
    :param caminho_inicial: Caminho inicial para salvar o documento.
    :return: Caminho final onde o documento foi salvo.
    """
    contador = 1
    caminho_final = caminho_inicial
    base, extensao = os.path.splitext(caminho_inicial)
    
    while True:
        try:
            doc.save(caminho_final)
            print(f"Documento salvo como: {caminho_final}")
            break  # Se o salvamento for bem-sucedido, saia do loop
        except PermissionError:
            # Se ocorrer um PermissionError, modificar o nome do arquivo
            caminho_final = f"{base}_{contador}{extensao}"
            contador += 1  # Aumenta o contador para o próximo nome
    return caminho_final

def main():
    # Caminho do documento
    caminho_documento = './BOP 000000000.000000-0 Celular crimes contra/Laudo prot - Celular extração e análise.docx'
    
    # Carregar o documento existente
    doc = Document(caminho_documento)
    
    # Caminho da imagem para transcrição
    caminho_imagem = './BOP 000000000.000000-0 Celular crimes contra/Figura 00 - Requisição.jpg'
    
    # Transcrever texto da imagem
    try:
        texto_extraido = transcrever_imagem(caminho_imagem)  # Chama a função para transcrição
    except FileNotFoundError:
        print(f"Erro: O arquivo de imagem não foi encontrado no caminho: {caminho_imagem}")
        return  # Encerra a execução se a imagem não for encontrada
    
    informacoes = processar_texto(texto_extraido)  # Extrai informações do texto transcrito

    # Exibir informações extraídas para verificação
    print("Informações extraídas:", informacoes)
    
    # Substituir texto entre chaves nos parágrafos
    for paragraph in doc.paragraphs:
        # Verifique se o parágrafo contém algum dos placeholders
        placeholders_presentes = any(f'{{{chave}}}' in paragraph.text for chave in informacoes.keys())
        if placeholders_presentes:
            print("Texto original do parágrafo:", paragraph.text)  # Imprime o parágrafo original
            substituir_placeholders_no_paragrafo(paragraph, informacoes)
            print("Texto modificado do parágrafo:", paragraph.text)  # Imprime o parágrafo modificado
    
    # Definir o caminho do novo documento
    caminho_novo_documento = './BOP 000000000.000000-0 Celular crimes contra/novo_documento.docx'
    
    # Salvar o documento modificado com tratamento de erro
    salvar_documento_com_tratamento(doc, caminho_novo_documento)

if __name__ == "__main__":
    main()

















# from docx import Document
# from funcoes_requisicao import processar_texto
# import os

# def substituir_placeholders_no_paragrafo(paragraph, informacoes):
#     """
#     Substitui os placeholders no parágrafo com os valores fornecidos.
    
#     :param paragraph: Objeto Paragraph do python-docx.
#     :param informacoes: Dicionário com chaves correspondentes aos placeholders e seus valores.
#     """
#     # Reconstrói o texto completo do parágrafo
#     full_text = ''.join(run.text for run in paragraph.runs)
    
#     # Realiza as substituições no texto completo
#     for chave, valor in informacoes.items():
#         placeholder = f'{{{chave}}}'
#         full_text = full_text.replace(placeholder, valor)
    
#     # Limpa os textos existentes nos runs
#     for run in paragraph.runs:
#         run.text = ''
    
#     # Adiciona um novo run com o texto substituído
#     paragraph.add_run(full_text)

# def salvar_documento_com_tratamento(doc, caminho_inicial):
#     """
#     Salva o documento, tratando erros de permissão adicionando um número ao nome do arquivo se necessário.
    
#     :param doc: Objeto Document do python-docx.
#     :param caminho_inicial: Caminho inicial para salvar o documento.
#     :return: Caminho final onde o documento foi salvo.
#     """
#     contador = 1
#     caminho_final = caminho_inicial
#     base, extensao = os.path.splitext(caminho_inicial)
    
#     while True:
#         try:
#             doc.save(caminho_final)
#             print(f"Documento salvo como: {caminho_final}")
#             break  # Se o salvamento for bem-sucedido, saia do loop
#         except PermissionError:
#             # Se ocorrer um PermissionError, modificar o nome do arquivo
#             caminho_final = f"{base}_{contador}{extensao}"
#             contador += 1  # Aumenta o contador para o próximo nome
#     return caminho_final

# def main():
#     # Caminho do documento
#     caminho_documento = './BOP 000000000.000000-0 Celular crimes contra/Laudo prot - Celular extração e análise.docx'
    
#     # Carregar o documento existente
#     doc = Document(caminho_documento)
    
#     # Texto simulado extraído para o exemplo (pode ser substituído pelo texto extraído via OCR)
#     texto_extraido = """
#     Atendimento a Nº 00001-2024-100059-1, datada de 02/10/2024, referente ao INQUERITO POR PORTARIA Nº 00001/2024.113864-9 
#     DIVISÃO DE COMBATE A CRIMES CONTRA DIREITOS INDIVIDUAIS POR MEIOS CIBERNÉTICOS, e assinado pela autoridade X.
#     """
    
#     # Processar o texto para extrair as informações
#     informacoes = processar_texto(texto_extraido)
    
#     # Exibir informações extraídas para verificação
#     print("Informações extraídas:", informacoes)
    
#     # Substituir texto entre chaves nos parágrafos
#     for paragraph in doc.paragraphs:
#         # Verifique se o parágrafo contém algum dos placeholders
#         placeholders_presentes = any(f'{{{chave}}}' in paragraph.text for chave in informacoes.keys())
#         if placeholders_presentes:
#             print("Texto original do parágrafo:", paragraph.text)  # Imprime o parágrafo original
#             substituir_placeholders_no_paragrafo(paragraph, informacoes)
#             print("Texto modificado do parágrafo:", paragraph.text)  # Imprime o parágrafo modificado
    
#     # Definir o caminho do novo documento
#     caminho_novo_documento = './BOP 000000000.000000-0 Celular crimes contra/novo_documento.docx'
    
#     # Salvar o documento modificado com tratamento de erro
#     salvar_documento_com_tratamento(doc, caminho_novo_documento)

# if __name__ == "__main__":
#     main()












