from docx import Document

def gerar_laudo(informacoes, caminho_arquivo):
    doc = Document()
    
    doc.add_heading('Laudo Pericial', 0)
    doc.add_paragraph(f"Nome: {informacoes['nome']}")
    doc.add_paragraph(f"Data: {informacoes['data']}")
    doc.add_paragraph(f"Protocolo: {informacoes['protocolo']}")
    doc.add_paragraph("Este é um laudo pericial gerado automaticamente.")
    
    doc.save(caminho_arquivo)
    print(f"Laudo salvo em: {caminho_arquivo}")

# Exemplo de uso da função
gerar_laudo(informacoes, "laudo_pericial.docx")
